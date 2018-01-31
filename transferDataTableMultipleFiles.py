#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Oct 16 10:03:55 2017

@author: Anthony Yorita
"""
import docx
import re
import openpyxl
# 2nd Script Gather All Data by Table
#***def transferDataTable(filename, NumberOfRows, wbLoad): 
    #****document = docx.Document(filename)
#document = docx.Document('/Users/Gesetze/Documents/Data_Science_Internship/Files_To_Transfer/17-077.docx')

def transferData(filename, NumberOfRows, wbLoad):
    document = docx.Document(filename)
    
        # Uncomment (remove # in front of line) below if you desire to know the 
        # number of tables in the word file
        # Data will be a list of rows represented as dictionaries or tuples
        # containing each row's data.
    data = []
    table_list = []
    allText = ''
    everyText = ''
    allTextList = []
    count = 0
    for table in document.tables: #document.tables are the tables objects
        #allText = ''
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            for cell in row.cells:
                #allText = allText + cell.text
                everyText = everyText + cell.text
                if 'Contact info' in cell.text or 'Contact Info' in cell.text:
                    count += 1
                if count == 1:
                    allTextList.append(cell.text)
                    for word in allTextList:
                        allText = allText + word
                    # print("\n--->allText", allText, "\n")
                    # print("\n", allText, "\n")
                    print(everyText)
                    table_list.append(everyText)
                    allText = ''
                    allTextList = []
                    count = 0
                    everyText = ''
                    #print(allText)
                else:
                    allTextList.append(cell.text)
                    #print(cell.text)
                #xsallText = allText + ' ' + cell.text
               # print(cell.text)
           # allText = allText + str(text)
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue
        
            # Construct a dictionary(or append a tuple) for this row, mapping
            # keys to values for this row
            
            row_data = tuple(text)
            #allText = allText + str(row_data)
            data.append(row_data)
        #table_list.append(allText)
    #print(allText)
    #print(allText)
   # print(len(table_list)) # prints only the first table
   # print(table_list[4])
    # all Regexes
    # ^ and $ aren't working because it means whatever the string starts with
    # not whether there is a ^WSP in the STRING!!!'
    
    # not from table
   # dateRegex = re.compile(r'\d\d-\d\d-\d\d|\d-\d-\d\d|\d\d-\d-\d\d|\d-\d\d-\d\d')
    dateRegex = re.compile(r'\d\d(-|/)?\d\d(-|/)?\d\d|\d(-|/)?\d(-|/)?\d\d|\d\d(-|/)?\d(-|/)?\d\d|\d(-|/)?\d\d(-|/)?\d\d')
    nibinRegex = re.compile(r'\d\d-\d\d\d')
    
    WSP_case_number_Regex = re.compile(r'(WSP Case Number)(.*)(Case Number)')
    Case_Number_regex = re.compile(r'(Case Number)(.*)(Exhibit Number)')
    Exhibit_regex = re.compile(r'(Exhibit Number)(.*)(Offense)')
    Offense_regex = re.compile(r'(Offense)(.*)(Evidence)')
    Evidence_regex = re.compile(r'(Evidence)(.*)(Date)')
    Date_regex = re.compile(r'(Date)(.*)(Location)')
    Location_regex = re.compile(r'(Location)(.*)(Agency)')
    Agency_regex = re.compile(r'(Agency)(.*)(Assigned Detective)')
    AD_regex = re.compile(r'(Assigned Detective)(.*)(Contact Info)')
    Contact_regex = re.compile(r'(Contact Info)(.*)(EVENT#)')

    
    # For finall method (everyTExt)
#    WSP_case_number_Regex = re.compile(r'(WSP case number)(.*?)(Case Number)')
#    Case_Number_regex = re.compile(r'(Case Number)(.*?)(Exhibit Number)')
#    Exhibit_regex = re.compile(r'(Exhibit Number)(.*?)(Offense)')
#    Offense_regex = re.compile(r'(Offense)(.*?)(Evidence)')
#    Evidence_regex = re.compile(r'(Evidence)(.*?)(Date)')
#    Date_regex = re.compile(r'(Date)(.*?)(Location)')
#    Location_regex = re.compile(r'(Location)(.*?)(Agency)')
#    Agency_regex = re.compile(r'(Agency)(.*?)(Assigned Detective)')
#    AD_regex = re.compile(r'(Assigned Detective)(.*?)(Contact info)')
#    Contact_regex = re.compile(r'(Contact info)(.*?)')
#    mo_case1 = Case_Number_regex.findall(allText)
#    print(mo_case1[0])
#    print(len(mo_case1))
#    
    #mo.group(0) prints all group. Starts at group(1), group(2)...
    #print(mo_Case.group(2))
    
    stringWord = ''
    
    #takes all paragraphs and put them into one string for data and NIBIN
    for line in document.paragraphs:
        if '_sre.SRE_Match' not in line.text:
            stringWord = stringWord + line.text
    mo = dateRegex.search(stringWord)
    moNibin = nibinRegex.search(stringWord)
    #print(mo.group())
    #print(moNibin.group())
    
    #  NumberOfRows = 2 #***** TAKE THIS OUT FOR FUNCTION!!!
    #leftSide = 'C' + str(NumberOfRows)
    #rightSide = 'L' + str(NumberOfRows + len(document.tables))
    #leftSideRegex = 'A' + str(NumberOfRows)
    #rightSideRegex = 'B' + str((NumberOfRows) + len(document.tables) - 2)
    
    #print(leftSideRegex)
    #print(rightSideRegex)
    
    # Adding table values to excel sheet
    # value is just for the column[x], it just goes to next
    # value in the sorted_data
    
    # WRITING TO EXCEL PORTION
    # Column A in excel will be date of investigation, B will be the NIBIN, 
    # C will be WSP case number,... and so on until L being Contact INfo
    # wbLoad is the excel workbook being loaded
    # sheet is the only sheet in the workbook, can specify more 
    wb = openpyxl.load_workbook(wbLoad)
    sheet = wb.get_sheet_by_name('Sheet1')
    
    """
    for i in range(len(document.tables) - 1):#range(2):
        mo_WSP = WSP_case_number_Regex.search(table_list[i])
        mo_Case = Case_Number_regex.search(table_list[i])
        mo_Exhibit = Exhibit_regex.search(table_list[i])
        mo_Offense = Offense_regex.search(table_list[i])
        mo_Date = Date_regex.search(table_list[i])
        mo_Location = Location_regex.search(table_list[i])
        mo_Agency = Agency_regex.search(table_list[i])
        mo_AD = AD_regex.search(table_list[i])
        mo_Contact = Contact_regex.search(table_list[i])
        print(mo_Case.group(2))
    """
    #print(sheet.cell(row = 1, column = 2))
    # to get cell do sheet.cell(row = x, column = y)
    # to get cell value do sheet.cell(row = x, column = y).value
    """
    value = 0
    for rowOfCellObjects in sheet[leftSide : rightSide]:
        for cellObj in rowOfCellObjects:
            #print(cellObj)
            #if value > len(sorted_data) - 1: 
            #    break
            #else: 
            #cellObj.value = sorted_data[value]
            value += 1
    """

    # print(table_list[0])
    for i in range(len(table_list)):
        # those two don't work because they are not in the table,
        # other code for mo and moNibin is above
       # mo = dateRegex.search(table_list[i])
        #moNibin = nibinRegex.search(table_list[i])
        mo_WSP = WSP_case_number_Regex.search(table_list[i])
        mo_Case = Case_Number_regex.search(table_list[i])
        mo_Exhibit = Exhibit_regex.search(table_list[i])
        mo_Offense = Offense_regex.search(table_list[i])
        mo_Evidence = Evidence_regex.search(table_list[i])
        mo_Date = Date_regex.search(table_list[i])
        mo_Location = Location_regex.search(table_list[i])
        mo_Agency = Agency_regex.search(table_list[i])
        mo_AD = AD_regex.search(table_list[i])
        mo_Contact = Contact_regex.search(table_list[i])

        # then input in
        regex_list = [mo, moNibin, mo_WSP, mo_Case, mo_Exhibit, mo_Offense,
                      mo_Evidence, mo_Date, mo_Location, mo_Agency, mo_AD,
                      mo_Contact]
        # Have to enter these manually since they are different than the rest
        # they don't have multiple groups in regex
        
        if (mo != None):
            sheet.cell(row= NumberOfRows + i, column= 1, value=mo.group())
        else :
            sheet.cell(row = NumberOfRows + i, column = 1, value = '')
        if (moNibin != None):
            sheet.cell(row= NumberOfRows + i, column=2, value=moNibin.group())
        else:
            sheet.cell(row = NumberOfRows + i, column = 2, value = '')
        
        columnNumber = 3
        for regex in range(2, len(regex_list)):
            # print(regex_list[regex])
            if regex_list[regex] != None and regex_list[regex]!= mo and regex_list[regex] != moNibin:
                value = regex_list[regex].group(2)
                if(regex == 3):
                    value = re.sub(r".*(Case Number)", "", value)
                sheet.cell(row = NumberOfRows + i, column = columnNumber, value = value)
            else:
                sheet.cell(row = NumberOfRows + i, column = columnNumber, value = '')

            columnNumber = columnNumber + 1
        sheet.cell(row= NumberOfRows + i, column= 13, value=filename)
        sheet.cell(row= NumberOfRows + i, column= 14, value='Table ' + str(i))
            #print(sheet.cell(row = 2, column = 1).value)
        
    
    # Directory format should be changed according to OS
    # The one below works for Macs    
    wb.save('KCNIBNRealTest3.xlsx') 
    tupe1 = (len(table_list), 'KCNIBNRealTest3.xlsx')
    return tupe1
        
