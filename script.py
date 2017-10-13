from docx import Document
import xlsxwriter
from collections import defaultdict

document = Document("C:\KCPAO\LEAD_Test_Document.docx") #use the os library to track through the 50 files in a directory
table_one = document.tables[0] #take out the first table
table_two = document.tables[1] #take out the second table
table_three = document.tables[2] #take out the third table
#table_four = document.tables[3] #take out the fourth table

def wordToList(table):
    final = {}
    for i, row in enumerate(table.rows):
        if (i != 0 and i != 1):
            text = (cell.text for cell in row.cells)
            text = list(text)
            final[text[0]] = text[1]
    return (final)

d1 = wordToList(table_one)
d2 = wordToList(table_two)
#d3 = wordToList(table_three)
#d4 = wordToList(table_four)

document_one_data = defaultdict(list)
for d in (d1, d2):#, d3, d4):
    for key, value in d.items():
        document_one_data[key].append(value)
print (document_one_data)

workbook = xlsxwriter.Workbook('test.xlsx')
worksheet = workbook.add_worksheet()
worksheet.set_column(0, 8, 30) # setting column widths


def writeToExcel(data, worksheet, row):

    for i, item in enumerate(document_one_data.values()):
        row = 1
        for j in range(len(item)):
            worksheet.write(row, i, item[j])
            row += 1

def writeHeader(worksheet):
    for i, item in enumerate(document_one_data.keys()):
        worksheet.write(0, i, item)

writeHeader(worksheet)
writeToExcel(document_one_data, worksheet, 1)
workbook.close()
