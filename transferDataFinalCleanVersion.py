#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  9 13:29:09 2017

@author: Anthony Yorita
"""

# Excel and Doc Word Transfer Together


import docx
import re
import openpyxl

# Load the file by using the docx module. 
# Can replace filename with an actual filename if desired
# Microsoft WORD PORTION !!!!
def transferData(filename, NumberOfRows, wbLoad): 
    document = docx.Document(filename)
    
    # Uncomment (remove # in front of line) below if you desire to know the 
    # number of tables in the word file
    # print(len(document.tables))
        
 
    # Data will be a list of rows represented as dictionaries or tuples
    # containing each row's data.
    data = []
    for table in document.tables: #document.tables are the tables objects
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue
        
            # Construct a dictionary(or append a tuple) for this row, mapping
            # keys to values for this row
            
            row_data = tuple(text)
            data.append(row_data)
    # Sorted_Data creates a list of WSP Case number, Case Number...
    # and so on in sorted order. So it comes out WSP, Case, Exhibit, Offense...        
    sorted_data = []
    for column in data:
        sorted_data.append(column[1])
        # column[0] contains the categorical variables 
        # column[1] contains the actual values inside the table
        # Uncomment below to see what is in the data
        # print('' + column[0] + ' : ' + column[1])
        
    
    # WRITING TO EXCEL PORTION
    # Column A in excel will be date of investigation, B will be the NIBIN, 
    # C will be WSP case number,... and so on until L being Contact INfo
    # wbLoad is the excel workbook being loaded
    # sheet is the only sheet in the workbook, can specify more 
    wb = openpyxl.load_workbook(wbLoad)
    sheet = wb.get_sheet_by_name('Sheet1')
    #(sheet['C2':'L13']) # gets a rectangular from C2 to L13
    # will go C2, d2, e2... unil l2, then back to C3
    
    
    # Regex Portion to find date and NIBIN Lead $
    # Regex is essentially using code to find a specific format (like CTRL Find)
    # dateRegex is set to find the files with the specific date XX/XX/XX
    # Can be changed to find any other types of dates
    # NibinRegex is set to find the NIBIN # with its specific format.
    # Can use '|' if the format can be 
    dateRegex = re.compile(r'\d\d-\d\d-\d\d|\d-\d-\d\d|\d\d-\d-\d\d|\d-\d\d-\d\d' )
    nibinRegex = re.compile(r'\d\d-\d\d\d')
    
    #The .search function actual finds the specific values in the document
    # or in this case paragraphs object. 
    # We use paragraphs[x].text to convert the paragraphs object to text
    # so the regex can find the specific values
    
    # Uncomment below to check the amount of paragraphs in document
    #print(len(document.paragraphs))
    
    #mo = dateRegex.search(document.paragraphs[0].text)
    #moNibin = nibinRegex.search(document.paragraphs[1].text)
    addUp = []
    stringWord = ''
    
    for line in document.paragraphs:
        if '_sre.SRE_Match' not in line.text:
            stringWord = stringWord + line.text
            #mo = dateRegex.search(line.text)
            #moNibin = nibinRegex.search(line.text)
    
    
    mo = dateRegex.search(stringWord)
    moNibin = nibinRegex.search(stringWord)
    #print(stringWord)
    #print(addUp)
    
    # Uncomment below to check to see if regex for date and NIBIN
    # searched correctly
    # print(mo.group())
    # print(moNibin.group())
    
    
    #moNibin = nibinRegex.search(document.paragraphs[1].text)
    #print(moNibin)
    # Uncomment below to check whether the regexe's are accurate or not
    #print('Date of case found: ' + mo.group())
    #print('Nibin Lead# found: ' + moNibin.group())
    
    # The sheet[value1 : value2] searches from cell value1 to value2
    # It makes a rectangular box if the cell for instance is B2 and D10
    # It will go B2, C2, D2 then back to B3, C3,.. and so on 
    # the leftSide and RightSide is for the values in the table
    # I had to separate the regexes because they aren't in the table 
    # in word. So if I had to incorporate it into one for loop, it would
    # be more complex since I would have to skip over certain columns
    leftSide = 'C' + str(NumberOfRows)
    rightSide = 'L' + str(NumberOfRows + len(document.tables))
    leftSideRegex = 'A' + str(NumberOfRows)
    rightSideRegex = 'B' + str((NumberOfRows) + len(document.tables) - 2)
    
    print(leftSideRegex)
    print(rightSideRegex)
    
    # Adding table values to excel sheet
    # value is just for the column[x], it just goes to next
    # value in the sorted_data
    value = 0
    for rowOfCellObjects in sheet[leftSide : rightSide]:
        for cellObj in rowOfCellObjects:
            if value > len(sorted_data) - 1: 
                break
            else: 
                cellObj.value = sorted_data[value]
                value += 1
    
    # Adding regex values to excel sheet
    # .group() just puts the entire regex in 
    # You can specific if you wish, which parts of a regex are in a certain
    # group
    
    for rowOfCellObjects in sheet[leftSideRegex : rightSideRegex]:
        for cellObj in rowOfCellObjects:
            if 'A' in cellObj.coordinate:
                cellObj.value = mo.group()
            #print (cellObj.coordinate) use .help(cellObj) to learn more about
            if 'B' in cellObj.coordinate:
                cellObj.value = moNibin.group()
                
    
    # Finally, you have to save the results in the excel file
    # You can save it into whichever you wish but make sure to change
    # tupe1's 2nd value to have what is being saved so that it doesn't
    # just overwrite the previous values each time it goes to another file
    
    #wb.save('Lead_TEST_DocumentVersion3.xlsx') # need to save
    wb.save('KCNIBN.xlsx')
    
    #tupe1 = (len(document.tables), 'Lead_TEST_DocumentVersion3.xlsx')
    tupe1 = (len(document.tables), 'KCNIBN.xlsx')
    # I return tupe1 since I need to know how many tables are in each word doc
    # This allows me to increment the rows properly. Also I return the
    # file the information is being saved in.
    return tupe1
